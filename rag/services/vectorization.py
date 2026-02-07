"""
Service de vectorisation de documents.

Responsabilités :
  - Extraction du texte brut (PDF, DOCX, XLSX, TXT, MD)
  - Découpage en chunks avec chevauchement

Le stockage des embeddings en pgvector est géré par la tâche Celery
(rag/tasks.py) qui appelle albert_client puis fait un bulk_create
sur DocumentChunk.

Porté depuis chatbot/services/vectorization.py — la partie ChromaDB
a été retirée (remplacée par l'ORM Django + pgvector).
"""
import os
import logging
from typing import List, Dict

import fitz  # PyMuPDF
import pdfplumber  # Extraction de tableaux PDF
from docx import Document as DocxDocument
import openpyxl
import tiktoken
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger('rag')


class VectorizationService:
    """Extraction et découpage de texte depuis des documents uploadés."""

    def __init__(self):
        # Tokenizer pour comptage précis (cl100k_base = GPT-4/embeddings modernes)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

        # Chunking basé sur tokens (512 tokens ≈ optimal pour BGE-M3)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,  # tokens (vs 1000 chars avant)
            chunk_overlap=100,  # tokens (vs 200 chars avant)
            length_function=self._count_tokens,  # Utilise tiktoken
            separators=[
                "\n\n=== TABLEAU",  # Préserver les tableaux ensemble
                "\n\n--- Page",     # Préserver les pages ensemble si possible
                "\n\n",             # Paragraphes
                "\n",               # Lignes
                ". ",               # Phrases
                " ",                # Mots
                ""
            ],
        )
        # Métadonnées extraites lors de l'extraction (pour stocker dans Document.metadata)
        self.extracted_metadata = {}

    def _count_tokens(self, text: str) -> int:
        """
        Compte le nombre de tokens dans un texte avec tiktoken.

        Args:
            text: Texte à analyser

        Returns:
            Nombre de tokens
        """
        return len(self.tokenizer.encode(text))

    # ------------------------------------------------------------------
    # Extraction de texte selon le format
    # ------------------------------------------------------------------

    def extract_text(self, file_path: str, filename: str) -> str:
        """
        Extrait le texte brut d'un fichier selon son extension.

        Formats supportés : .pdf .docx .xlsx .txt .md

        Note: Les métadonnées extraites (stats XLSX, nombre de pages PDF, etc.)
        sont stockées dans self.extracted_metadata pour être sauvegardées
        dans Document.metadata par la tâche Celery.
        """
        ext = os.path.splitext(filename)[1].lower()

        # Réinitialiser les métadonnées pour chaque extraction
        self.extracted_metadata = {'format': ext[1:]}  # 'pdf', 'xlsx', etc.

        extractors = {
            '.pdf':  self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_xlsx,
            '.txt':  self._extract_text,
            '.md':   self._extract_text,
        }

        extractor = extractors.get(ext)
        if not extractor:
            raise ValueError(f"Format non supporté : {ext}")

        try:
            text = extractor(file_path)
            logger.info(f"Texte extrait de {filename} : {len(text)} caractères")
            return text
        except Exception as e:
            logger.error(f"Erreur extraction {filename} : {e}")
            raise

    def _extract_pdf(self, file_path: str) -> str:
        """
        PDF → texte avec tableaux (PyMuPDF + pdfplumber).

        Stratégie:
        - PyMuPDF pour le texte général (rapide, fiable)
        - pdfplumber pour détecter et extraire les tableaux (précis pour tableaux)
        - Fusion intelligente : texte + tableaux formatés
        """
        # 1. Extraction texte avec PyMuPDF
        doc_fitz = fitz.open(file_path)
        pages_content = []
        table_count = 0

        # 2. Extraction tableaux avec pdfplumber
        with pdfplumber.open(file_path) as pdf_plumber:
            for page_num, (page_fitz, page_plumber) in enumerate(
                zip(doc_fitz, pdf_plumber.pages), start=1
            ):
                page_text = f"\n\n--- Page {page_num} ---\n\n"

                # Texte de la page (PyMuPDF)
                text_content = page_fitz.get_text()

                # Tableaux de la page (pdfplumber)
                tables = page_plumber.extract_tables()

                if tables:
                    # Si la page contient des tableaux, les formater proprement
                    page_text += f"📊 Cette page contient {len(tables)} tableau(x)\n\n"
                    page_text += text_content

                    # Ajouter chaque tableau formaté
                    for table_idx, table in enumerate(tables, start=1):
                        table_count += 1
                        page_text += f"\n\n=== TABLEAU {table_idx} (Page {page_num}) ===\n"
                        page_text += self._format_table(table)
                        page_text += "\n=== FIN TABLEAU ===\n\n"
                else:
                    # Pas de tableau, juste le texte
                    page_text += text_content

                pages_content.append(page_text)

        # 3. Stocker les métadonnées PDF
        self.extracted_metadata.update({
            'page_count': len(doc_fitz),
            'has_toc': len(doc_fitz.get_toc()) > 0,
            'has_tables': table_count > 0,
            'table_count': table_count
        })
        logger.info(
            f"PDF metadata: {len(doc_fitz)} page(s), {table_count} tableau(x) extrait(s)"
        )

        doc_fitz.close()
        return "".join(pages_content)

    def _format_table(self, table: List[List]) -> str:
        """
        Formate un tableau extrait par pdfplumber en texte lisible.

        Args:
            table: Liste de listes (lignes × colonnes)

        Returns:
            Tableau formaté en texte avec séparateurs
        """
        if not table or len(table) == 0:
            return "(Tableau vide)"

        # Nettoyer les cellules vides (None → "")
        cleaned_table = [
            [str(cell).strip() if cell else "" for cell in row]
            for row in table
        ]

        # Calculer la largeur maximale de chaque colonne
        col_widths = []
        if cleaned_table:
            num_cols = max(len(row) for row in cleaned_table)
            for col_idx in range(num_cols):
                max_width = max(
                    len(row[col_idx]) if col_idx < len(row) else 0
                    for row in cleaned_table
                )
                col_widths.append(min(max_width, 40))  # Max 40 chars par colonne

        # Formater chaque ligne
        formatted_rows = []
        for row_idx, row in enumerate(cleaned_table):
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                width = col_widths[col_idx] if col_idx < len(col_widths) else 20
                # Tronquer si trop long
                cell_text = cell[:width].ljust(width)
                formatted_cells.append(cell_text)
            formatted_rows.append(" | ".join(formatted_cells))

            # Ligne de séparation après l'en-tête (première ligne)
            if row_idx == 0 and len(cleaned_table) > 1:
                separator = "-+-".join(["-" * w for w in col_widths[:len(row)]])
                formatted_rows.append(separator)

        return "\n".join(formatted_rows)

    def _extract_docx(self, file_path: str) -> str:
        """
        DOCX → texte avec tableaux (python-docx).

        Extrait les paragraphes ET les tableaux dans l'ordre d'apparition.
        Préserve la structure hiérarchique du document.
        """
        doc = DocxDocument(file_path)
        content_parts = []
        table_count = 0

        # Parcourir tous les éléments du document dans l'ordre
        # Note: doc.element.body permet d'accéder aux éléments dans l'ordre
        for element in doc.element.body:
            # Paragraphe
            if element.tag.endswith('p'):
                # Trouver le paragraphe correspondant
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:  # Ignorer les paragraphes vides
                            content_parts.append(text)
                        break

            # Tableau
            elif element.tag.endswith('tbl'):
                # Trouver le tableau correspondant
                for table in doc.tables:
                    if table._element == element:
                        table_count += 1
                        content_parts.append(
                            f"\n\n=== TABLEAU {table_count} ===\n"
                        )
                        content_parts.append(self._format_docx_table(table))
                        content_parts.append("\n=== FIN TABLEAU ===\n\n")
                        break

        # Stocker les métadonnées DOCX
        self.extracted_metadata.update({
            'paragraph_count': len(doc.paragraphs),
            'has_tables': table_count > 0,
            'table_count': table_count
        })
        logger.info(
            f"DOCX metadata: {len(doc.paragraphs)} paragraphe(s), "
            f"{table_count} tableau(x)"
        )

        return "\n\n".join(content_parts)

    def _format_docx_table(self, table) -> str:
        """
        Formate un tableau DOCX en texte lisible.

        Args:
            table: Objet Table de python-docx

        Returns:
            Tableau formaté en texte avec séparateurs
        """
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            return "(Tableau vide)"

        # Calculer la largeur de chaque colonne
        num_cols = max(len(row) for row in rows_data) if rows_data else 0
        col_widths = []
        for col_idx in range(num_cols):
            max_width = max(
                len(row[col_idx]) if col_idx < len(row) else 0
                for row in rows_data
            )
            col_widths.append(min(max_width, 40))  # Max 40 chars

        # Formater chaque ligne
        formatted_rows = []
        for row_idx, row in enumerate(rows_data):
            formatted_cells = []
            for col_idx, cell in enumerate(row):
                width = col_widths[col_idx] if col_idx < len(col_widths) else 20
                cell_text = cell[:width].ljust(width)
                formatted_cells.append(cell_text)
            formatted_rows.append(" | ".join(formatted_cells))

            # Séparateur après l'en-tête
            if row_idx == 0 and len(rows_data) > 1:
                separator = "-+-".join(["-" * w for w in col_widths[:len(row)]])
                formatted_rows.append(separator)

        return "\n".join(formatted_rows)

    def _extract_xlsx(self, file_path: str) -> str:
        """XLSX → texte avec noms de feuilles et métadonnées statistiques."""
        wb = openpyxl.load_workbook(file_path)
        sheets = []

        # Métadonnées globales du fichier
        xlsx_metadata = {
            'sheets': [],
            'total_data_rows': 0,
            'total_rows': 0
        }

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []

            # Extraire toutes les lignes
            all_rows = list(sheet.iter_rows(values_only=True))

            # Générer les métadonnées statistiques
            total_rows = len(all_rows)
            data_rows = total_rows - 1 if total_rows > 0 else 0  # Sans l'en-tête

            # Stocker les métadonnées de cette feuille
            sheet_metadata = {
                'name': sheet_name,
                'total_rows': total_rows,
                'data_rows': data_rows,
                'columns': []
            }

            # En-tête de feuille avec statistiques
            header = f"\n\n--- Feuille : {sheet_name} ---\n"
            header += f"STATISTIQUES : {data_rows} lignes de données, {total_rows} lignes au total"

            # Si on a un en-tête, l'afficher
            if all_rows and len(all_rows) > 0:
                first_row = all_rows[0]
                columns = [str(c) for c in first_row if c is not None]
                if columns:
                    sheet_metadata['columns'] = columns
                    header += f", {len(columns)} colonnes\n"
                    header += f"COLONNES : {' | '.join(columns)}\n"

            header += "\n"

            # Convertir toutes les lignes en texte
            for row in all_rows:
                rows.append(" | ".join(str(c) for c in row if c is not None))

            sheets.append(header + "\n".join(rows))

            # Accumuler les statistiques globales
            xlsx_metadata['sheets'].append(sheet_metadata)
            xlsx_metadata['total_data_rows'] += data_rows
            xlsx_metadata['total_rows'] += total_rows

        # Stocker les métadonnées pour utilisation dans tasks.py
        self.extracted_metadata.update(xlsx_metadata)
        logger.info(
            f"XLSX metadata: {len(xlsx_metadata['sheets'])} feuille(s), "
            f"{xlsx_metadata['total_data_rows']} lignes de données"
        )

        return "".join(sheets)

    def _extract_text(self, file_path: str) -> str:
        """TXT / MD → lecture brute UTF-8."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

    # ------------------------------------------------------------------
    # Découpage en chunks
    # ------------------------------------------------------------------

    def chunk_text(self, text: str, filename: str, min_chunk_size: int = 50) -> List[Dict]:
        """
        Découpe un texte en chunks avec métadonnées complètes.

        Utilise un chunking basé sur tokens (512 tokens par chunk, 100 tokens d'overlap)
        pour une meilleure compatibilité avec les modèles d'embeddings (BGE-M3).

        Filtre les chunks trop courts (pages vides, headers seuls, etc.)
        pour éviter de polluer la base vectorielle.

        Args:
            text: Texte à découper
            filename: Nom du fichier source (pour traçabilité)
            min_chunk_size: Taille minimale d'un chunk (défaut: 50 caractères)

        Returns:
            [
                {
                    "content": "texte du chunk",
                    "metadata": {
                        "source": "filename.pdf",
                        "chunk_index": 0,
                        "original_index": 0,
                        "char_count": 950,
                        "token_count": 245
                    }
                },
                ...
            ]
        """
        raw_chunks = self.text_splitter.split_text(text)

        # Filtrer les chunks trop courts en conservant les métadonnées
        filtered_chunks = []
        original_index = 0
        total_tokens = 0

        for chunk in raw_chunks:
            if len(chunk.strip()) >= min_chunk_size:
                token_count = self._count_tokens(chunk)
                total_tokens += token_count

                filtered_chunks.append({
                    "content": chunk,
                    "metadata": {
                        "source": filename,
                        "chunk_index": len(filtered_chunks),
                        "original_index": original_index,
                        "char_count": len(chunk),
                        "token_count": token_count  # Nouveau: compte de tokens
                    }
                })
            original_index += 1

        if len(filtered_chunks) < len(raw_chunks):
            logger.info(
                f"Filtré {len(raw_chunks) - len(filtered_chunks)} chunks vides "
                f"({len(filtered_chunks)} chunks conservés)"
            )

        # Statistiques de chunking
        if filtered_chunks:
            avg_tokens = total_tokens / len(filtered_chunks)
            logger.info(
                f"Chunking: {len(filtered_chunks)} chunks, "
                f"moyenne {avg_tokens:.0f} tokens/chunk, "
                f"total {total_tokens} tokens"
            )

        return filtered_chunks
